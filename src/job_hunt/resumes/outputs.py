"""Create readable application DOCX files and layout-preserving resume PDFs."""

from __future__ import annotations

import base64
import os
import shutil
import subprocess
import tempfile
from datetime import date
from pathlib import Path
from typing import Any, Mapping, Sequence


class DocumentOutputError(RuntimeError):
    """Raised when a requested document format cannot be created safely."""


def _require_python_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Mm, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - dependency boundary
        raise RuntimeError("python-docx is required for cover-letter generation.") from exc
    return Document, WD_ALIGN_PARAGRAPH, qn, Inches, Mm, Pt, RGBColor


def _set_run_font(run, *, qn, Pt, RGBColor, size: float, color: str, bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bool(bold)


def build_cover_letter_docx(
    output_path: Path,
    *,
    identity: Mapping[str, str],
    posting: Mapping[str, Any],
    paragraphs: Sequence[str],
    generated_on: date,
) -> Path:
    """Build a compact A4 application letter without exposing identity to an LLM."""

    Document, WD_ALIGN_PARAGRAPH, qn, Inches, Mm, Pt, RGBColor = _require_python_docx()
    clean_paragraphs = [" ".join(str(value or "").split()) for value in paragraphs]
    clean_paragraphs = [value for value in clean_paragraphs if value]
    if not 2 <= len(clean_paragraphs) <= 5:
        raise DocumentOutputError("A cover letter requires two to five verified paragraphs.")
    if any("{{" in value or "[NAME]" in value.upper() for value in clean_paragraphs):
        raise DocumentOutputError("The cover letter contains an unresolved placeholder.")

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    # standard_business_brief with an A4 application-letter geometry override.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.85)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    name = " ".join(str(identity.get("name") or "Candidate").split())
    contact = " ".join(str(identity.get("contact_line") or "").split())
    name_paragraph = document.add_paragraph()
    name_paragraph.paragraph_format.space_after = Pt(2)
    name_run = name_paragraph.add_run(name)
    _set_run_font(
        name_run,
        qn=qn,
        Pt=Pt,
        RGBColor=RGBColor,
        size=16,
        color="0B2545",
        bold=True,
    )
    if contact:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.paragraph_format.space_after = Pt(16)
        contact_run = contact_paragraph.add_run(contact)
        _set_run_font(
            contact_run,
            qn=qn,
            Pt=Pt,
            RGBColor=RGBColor,
            size=9.5,
            color="52606D",
        )

    date_paragraph = document.add_paragraph(generated_on.strftime("%d %B %Y"))
    date_paragraph.paragraph_format.space_after = Pt(14)

    company = " ".join(str(posting.get("company") or "Hiring Team").split())
    location = " ".join(str(posting.get("location") or "").split())
    recipient = document.add_paragraph()
    recipient.paragraph_format.space_after = Pt(14)
    recipient.add_run("Hiring Team\n").bold = True
    recipient.add_run(company)
    if location:
        recipient.add_run(f"\n{location}")

    salutation = document.add_paragraph("Dear Hiring Team,")
    salutation.paragraph_format.space_after = Pt(10)

    for value in clean_paragraphs:
        paragraph = document.add_paragraph(value)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(9)
        paragraph.paragraph_format.line_spacing = 1.10

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(8)
    closing.paragraph_format.space_after = Pt(0)
    closing.add_run("Sincerely,\n")
    closing_name = closing.add_run(name)
    closing_name.bold = True

    pending = output_path.with_suffix(".pending.docx")
    document.save(pending)
    pending.replace(output_path)
    return output_path


def build_job_description_docx(
    output_path: Path,
    *,
    posting: Mapping[str, Any],
    description: str,
    completeness: str,
    description_source: str,
    capture_warning: str = "",
) -> Path:
    """Build a human-readable job-description record beside application documents."""

    Document, _, qn, Inches, Mm, Pt, RGBColor = _require_python_docx()
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.7)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(str(posting.get("title") or "Job description").strip())
    _set_run_font(
        run,
        qn=qn,
        Pt=Pt,
        RGBColor=RGBColor,
        size=18,
        color="0B2545",
        bold=True,
    )
    company = " ".join(str(posting.get("company") or "").split())
    if company:
        company_line = document.add_paragraph()
        company_line.paragraph_format.space_after = Pt(12)
        company_run = company_line.add_run(company)
        _set_run_font(
            company_run,
            qn=qn,
            Pt=Pt,
            RGBColor=RGBColor,
            size=11,
            color="276B63",
            bold=True,
        )

    metadata = (
        ("Location", posting.get("location")),
        ("Experience", posting.get("experience_text")),
        ("Employment", posting.get("employment_type")),
        ("Workplace", posting.get("workplace_type")),
        ("Requisition", posting.get("requisition_id")),
        ("Published", posting.get("published_at")),
        ("Official URL", posting.get("official_url")),
        ("Capture quality", completeness.replace("_", " ").title()),
        ("Capture source", description_source.replace("_", " ").title()),
    )
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        text = " ".join(str(value or "").split())
        if not text:
            continue
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = text
        cells[0].paragraphs[0].runs[0].bold = True

    if capture_warning:
        warning = document.add_paragraph()
        warning.paragraph_format.space_before = Pt(10)
        warning.paragraph_format.space_after = Pt(8)
        warning_run = warning.add_run(f"Review note: {capture_warning.strip()}")
        _set_run_font(
            warning_run,
            qn=qn,
            Pt=Pt,
            RGBColor=RGBColor,
            size=9.5,
            color="8A5A00",
            bold=True,
        )

    document.add_heading("Job description", level=1)
    clean_description = str(description or "").strip()
    if not clean_description:
        document.add_paragraph("The public source did not provide description text.")
    else:
        for raw_line in clean_description.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:].strip(), level=2)
            elif line.startswith(("- ", "• ", "* ")):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)

    footer = section.footer.paragraphs[0]
    footer.alignment = 1
    footer_run = footer.add_run(
        "Private application evidence · captured from the public official job source"
    )
    _set_run_font(
        footer_run,
        qn=qn,
        Pt=Pt,
        RGBColor=RGBColor,
        size=8,
        color="6B7280",
    )

    pending = output_path.with_suffix(".pending.docx")
    document.save(pending)
    pending.replace(output_path)
    return output_path


def _verify_pdf(path: Path) -> Path:
    path = Path(path).resolve()
    if not path.is_file() or path.stat().st_size < 1000:
        raise DocumentOutputError("The requested PDF was not created correctly.")
    with path.open("rb") as handle:
        if handle.read(5) != b"%PDF-":
            raise DocumentOutputError("The generated file is not a valid PDF.")
    return path


def _convert_with_soffice(executable: str, input_path: Path, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(
        prefix="job_hunt_pdf_", dir=output_path.parent
    ) as temporary:
        temporary_path = Path(temporary)
        completed = subprocess.run(
            [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temporary_path),
                str(input_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        produced = temporary_path / f"{input_path.stem}.pdf"
        if completed.returncode != 0 or not produced.is_file():
            raise DocumentOutputError("LibreOffice could not convert the tailored resume to PDF.")
        produced.replace(output_path)
    return _verify_pdf(output_path)


def _convert_with_word(input_path: Path, output_path: Path) -> Path:
    powershell = shutil.which("powershell.exe") or shutil.which("powershell")
    if not powershell:
        raise DocumentOutputError("Microsoft Word PDF export is unavailable.")
    script = r"""
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {
    $inputPath = $env:JOB_HUNT_PDF_INPUT
    $outputPath = $env:JOB_HUNT_PDF_OUTPUT
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Open($inputPath, $false, $true)
    $document.ExportAsFixedFormat($outputPath, 17)
}
finally {
    if ($null -ne $document) { $document.Close($false) }
    if ($null -ne $word) { $word.Quit() }
}
""".strip()
    encoded = base64.b64encode(script.encode("utf-16le")).decode("ascii")
    environment = dict(os.environ)
    environment["JOB_HUNT_PDF_INPUT"] = str(input_path)
    environment["JOB_HUNT_PDF_OUTPUT"] = str(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run(
        [powershell, "-NoProfile", "-NonInteractive", "-EncodedCommand", encoded],
        check=False,
        capture_output=True,
        text=True,
        timeout=120,
        env=environment,
    )
    if completed.returncode != 0:
        raise DocumentOutputError("Microsoft Word could not export the tailored resume to PDF.")
    return _verify_pdf(output_path)


def convert_docx_to_pdf(input_path: Path, output_path: Path) -> Path:
    """Convert a generated DOCX without saving or changing its source file."""

    input_path = Path(input_path).resolve()
    output_path = Path(output_path).resolve()
    if not input_path.is_file() or input_path.suffix.casefold() != ".docx":
        raise FileNotFoundError("The generated DOCX for PDF conversion is unavailable.")
    if output_path.suffix.casefold() != ".pdf":
        raise ValueError("The PDF output path must end in .pdf.")
    configured = os.environ.get("JOB_HUNT_SOFFICE_PATH", "").strip()
    executable = configured or shutil.which("soffice") or shutil.which("libreoffice")
    if executable:
        return _convert_with_soffice(executable, input_path, output_path)
    if os.name == "nt":
        return _convert_with_word(input_path, output_path)
    raise DocumentOutputError(
        "PDF generation requires LibreOffice on the deployed server. "
        "Install it or set JOB_HUNT_SOFFICE_PATH."
    )
